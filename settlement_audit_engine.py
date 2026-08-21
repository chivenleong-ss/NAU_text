"""Deterministic audit calculations and lightweight document normalization.

The demo deliberately keeps numeric conclusions in code. A language model can
later help with semantic extraction and explanations, but it must not decide
quantities, rates, amounts, or the final reduction amount.
"""

from __future__ import annotations

import csv
import io
import json
import re
import zipfile
from copy import deepcopy
from pathlib import Path
from typing import Any


DEFAULT_THRESHOLDS = {
    "quantity_deviation_pct": 0.05,
    "unit_price_deviation_pct": 0.10,
    "market_deviation_pct": 0.10,
    "duplicate_overlap_pct": 0.95,
}


def _number(value: Any, default: float = 0.0) -> float:
    if value is None or value == "":
        return default
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).replace(",", "").replace("，", "").replace("¥", "").replace("元", "")
    match = re.search(r"-?\d+(?:\.\d+)?", text)
    return float(match.group(0)) if match else default


def _money(value: float) -> int:
    return int(round(value))


def _pct(value: float, reference: float) -> float | None:
    if abs(reference) < 1e-9:
        return None
    return round((value - reference) / reference, 4)


def _amount(record: dict[str, Any]) -> float:
    if "amount" in record:
        return _number(record["amount"])
    return _number(record.get("quantity")) * _number(record.get("unit_price"))


def _risk(level: str) -> str:
    return {"high": "高风险", "medium": "中风险", "low": "低风险"}.get(level, level)


def compare_three_tables(data: dict[str, Any]) -> list[dict[str, Any]]:
    """Compare bid, contract, change, actual and settlement facts per item."""
    rows: list[dict[str, Any]] = []
    for item in data.get("line_items", []):
        bid = item.get("bid", {})
        contract = item.get("contract", {})
        change = item.get("change", {})
        actual = item.get("actual", {})
        settlement = item.get("settlement", {})
        contract_qty = _number(contract.get("quantity"))
        actual_qty = _number(actual.get("quantity"))
        settlement_qty = _number(settlement.get("quantity"))
        contract_price = _number(contract.get("unit_price"))
        settlement_price = _number(settlement.get("unit_price"))
        market_price = _number(item.get("market_unit_price"))
        rows.append({
            "item_code": item.get("item_code", ""),
            "name": item.get("name", ""),
            "unit": item.get("unit", ""),
            "location": item.get("location", ""),
            "bid": {"quantity": _number(bid.get("quantity")), "unit_price": _number(bid.get("unit_price")), "amount": _money(_amount(bid))},
            "contract": {"quantity": contract_qty, "unit_price": contract_price, "amount": _money(_amount(contract))},
            "change": {"quantity": _number(change.get("quantity")), "unit_price": _number(change.get("unit_price")), "amount": _money(_amount(change)), "id": change.get("id", ""), "approved": change.get("approved", True)},
            "actual": {"quantity": actual_qty, "unit_price": _number(actual.get("unit_price")), "amount": _money(_amount(actual))},
            "settlement": {"quantity": settlement_qty, "unit_price": settlement_price, "amount": _money(_amount(settlement))},
            "market_unit_price": market_price or None,
            "quantity_deviation_settlement_vs_actual": _pct(settlement_qty, actual_qty),
            "quantity_deviation_settlement_vs_contract": _pct(settlement_qty, contract_qty),
            "unit_price_deviation_settlement_vs_contract": _pct(settlement_price, contract_price),
            "unit_price_deviation_settlement_vs_market": _pct(settlement_price, market_price),
            "evidence": deepcopy(item.get("evidence", [])),
        })
    return rows


def _issue(issue_id: str, phase: str, category: str, title: str, level: str,
           amount: float, rule: str, judgement: str, suggestion: str,
           item: dict[str, Any], status: str = "待主审复核") -> dict[str, Any]:
    return {
        "id": issue_id,
        "phase": phase,
        "category": category,
        "title": title,
        "risk": _risk(level),
        "amount": _money(amount),
        "rule": rule,
        "judgement": judgement,
        "suggestion": suggestion,
        "item_code": item.get("item_code", ""),
        "location": item.get("location", ""),
        "source_objects": [x.get("document", "") for x in item.get("evidence", [])],
        "evidence": deepcopy(item.get("evidence", [])),
        "expert_status": "待分派",
        "chief_status": status,
    }


def detect_issues(data: dict[str, Any], comparisons: list[dict[str, Any]]) -> list[dict[str, Any]]:
    thresholds = dict(DEFAULT_THRESHOLDS)
    thresholds.update(data.get("meta", {}).get("thresholds", {}))
    items = {x.get("item_code"): x for x in data.get("line_items", [])}
    rows = {x.get("item_code"): x for x in comparisons}
    issues: list[dict[str, Any]] = []

    dirt = items.get("010101001")
    if dirt:
        change = dirt.get("change", {})
        contract_qty = _number(dirt.get("contract", {}).get("quantity"))
        change_qty = _number(change.get("quantity"))
        overlap = min(contract_qty, change_qty) / contract_qty if contract_qty else 0
        if change.get("relation") == "overlap" and overlap >= thresholds["duplicate_overlap_pct"]:
            issues.append(_issue("post-quantity-001", "事后", "工程量", "土方开挖重复计费", "high", change_qty * _number(dirt.get("contract", {}).get("unit_price")),
                f"同一清单编码、部位的合同与变更工程量重合率 {overlap:.0%} >= {thresholds['duplicate_overlap_pct']:.0%}",
                "结算书与 BG-018 均计取地下车库土方开挖，合同清单扣减关系尚未证明，存在重复计价。",
                "补充变更台账、现场收方记录和合同清单扣减记录；在核验前暂按重合工程量控制结算。", dirt))

    change_item = items.get("ZQ-023")
    if change_item and not change_item.get("change", {}).get("approved", True):
        change = change_item.get("change", {})
        issues.append(_issue("post-change-001", "事中", "变更签证", "变更 ZQ-023 审批链不完整", "high", _amount(change),
            "结算引用的变更记录 approval=false，且审批链缺少合同约定岗位",
            "ZQ-023 已进入结算申报，但项目商务经理签认缺失，变更计价效力存在程序性风险。",
            "冻结该项支付，补齐授权审批或转为现场核验事项。", change_item))

    temp = items.get("040801001")
    if temp and _number(temp.get("actual", {}).get("quantity")) == 0 and _number(temp.get("settlement", {}).get("quantity")) > 0:
        issues.append(_issue("post-actual-001", "事后", "实际完成量", "临时设施虚列子目", "high", _amount(temp.get("settlement", {})),
            "竣工实际完成量 = 0，结算工程量 > 0，且未找到移交证据",
            "结算书申报临时设施，但竣工移交资料未见实体移交记录，真实性不足。",
            "补充实体移交、拆除或现场核验记录；证据不足时不予计入结算。", temp))

    concrete = rows.get("010501002")
    concrete_item = items.get("010501002")
    if concrete and concrete_item and abs(concrete.get("unit_price_deviation_settlement_vs_contract") or 0) > thresholds["unit_price_deviation_pct"]:
        delta = (concrete["settlement"]["unit_price"] - concrete["contract"]["unit_price"]) * concrete["settlement"]["quantity"]
        issues.append(_issue("post-price-001", "事后", "单价", "混凝土结算单价偏离合同", "medium", delta,
            f"结算单价较合同单价偏离 {abs(concrete['unit_price_deviation_settlement_vs_contract']):.2%} > {thresholds['unit_price_deviation_pct']:.0%}",
            "结算单价高于合同约定，且未在当前资料包中发现完整调价依据。",
            "核对调价条款、审批单、供应商报价和同期市场参考价后再认定。", concrete_item))

    steel = items.get("010515001")
    leakage = steel.get("leakage", {}) if steel else {}
    if leakage.get("kind") == "material_loss":
        base = _number(leakage.get("base_quantity"))
        recorded = _number(leakage.get("recorded_consumption"))
        target = _number(leakage.get("target_loss_rate"))
        loss_rate = (recorded - base) / base if base else 0
        excess_amount = max(0, recorded - base * (1 + target)) * _number(leakage.get("unit_price"))
        if loss_rate > target:
            issues.append(_issue("post-leakage-001", "事后", "跑冒滴漏", "钢筋材料损耗率异常", "medium", excess_amount,
                f"实际损耗率 {loss_rate:.2%} > 目标损耗率 {target:.2%}",
                "材料台账领用量超过理论消耗与目标损耗允许范围，超额部分缺少专项说明。",
                "勾稽领用、退库、盘点和施工日志，补充损耗审批或核减超额材料成本。", steel))

    machine = items.get("021001004")
    machine_leakage = machine.get("leakage", {}) if machine else {}
    if machine_leakage.get("kind") == "machine_mismatch":
        unmatched = max(0, _number(machine_leakage.get("claimed_hours")) - _number(machine_leakage.get("matched_hours")))
        amount = unmatched / 8 * _number(machine_leakage.get("unit_price"))
        if unmatched:
            issues.append(_issue("post-leakage-002", "事后", "跑冒滴漏", "挖掘机台班与施工日志不匹配", "medium", amount,
                f"申报台班中有 {unmatched:.0f} 小时缺少施工日志匹配记录",
                "机械台班结算记录与施工日志存在连续缺口，无法证明对应设备实际作业。",
                "补充设备进退场、燃油、作业量和施工日志；证据不足部分暂不计量。", machine))

    if not data.get("meta", {}).get("data_source"):
        return issues

    existing_ids = {issue["id"] for issue in issues}
    for row in comparisons:
        item = items.get(row.get("item_code"), {"evidence": row.get("evidence", [])})
        code = row.get("item_code", "")
        title_name = row.get("name", code)
        settlement_amount = row.get("settlement", {}).get("amount", 0)
        actual_amount = row.get("actual", {}).get("amount", 0)
        quantity_deviation = row.get("quantity_deviation_settlement_vs_actual")
        if (
            quantity_deviation is not None
            and abs(quantity_deviation) > thresholds["quantity_deviation_pct"]
            and f"generic-quantity-{code}" not in existing_ids
        ):
            excess = max(0, settlement_amount - actual_amount)
            issues.append(_issue(
                f"generic-quantity-{code}",
                "事后",
                "工程量",
                f"{title_name}结算工程量偏离实际完成量",
                "high" if abs(quantity_deviation) >= 0.2 else "medium",
                excess,
                f"结算工程量较实际完成量偏离 {abs(quantity_deviation):.2%} > {thresholds['quantity_deviation_pct']:.0%}",
                "上传资料中的结算工程量与实际完成工程量不一致，超过阈值。",
                "核对竣工验收、现场收方和结算申报明细，按实际完成量复核。",
                item,
            ))
        if (
            row.get("actual", {}).get("quantity") == 0
            and row.get("settlement", {}).get("quantity", 0) > 0
            and f"generic-empty-actual-{code}" not in existing_ids
        ):
            issues.append(_issue(
                f"generic-empty-actual-{code}",
                "事后",
                "实际完成量",
                f"{title_name}结算有量但实际完成量为零",
                "high",
                settlement_amount,
                "实际完成工程量 = 0，结算工程量 > 0",
                "上传资料显示该子目未形成实际完成量，但已进入结算申报。",
                "补充验收或移交证据；证据不足时不予计入结算。",
                item,
            ))
        price_deviation = row.get("unit_price_deviation_settlement_vs_contract")
        if (
            price_deviation is not None
            and abs(price_deviation) > thresholds["unit_price_deviation_pct"]
            and f"generic-price-{code}" not in existing_ids
        ):
            quantity = row.get("settlement", {}).get("quantity", 0)
            contract_price = row.get("contract", {}).get("unit_price", 0)
            settlement_price = row.get("settlement", {}).get("unit_price", 0)
            issues.append(_issue(
                f"generic-price-{code}",
                "事后",
                "单价",
                f"{title_name}结算单价偏离合同",
                "medium",
                max(0, settlement_price - contract_price) * quantity,
                f"结算单价较合同单价偏离 {abs(price_deviation):.2%} > {thresholds['unit_price_deviation_pct']:.0%}",
                "上传资料中的结算单价偏离合同约定单价，超过规则阈值。",
                "核对调价条款、审批记录和市场参考价后再确认。",
                item,
            ))

    return issues


def _component_item(record: dict[str, Any], code: str, location: str) -> dict[str, Any]:
    return {
        "item_code": code,
        "location": location,
        "evidence": deepcopy(record.get("evidence", [])),
    }


def aggregate_cost_components(data: dict[str, Any]) -> dict[str, Any]:
    """Reconcile contract, changes, claims, rewards, penalties and fees."""
    project = data.get("project", {})
    financials = project.get("financials", {})
    components = project.get("cost_components", {})
    claims = components.get("claims", [])
    rewards_penalties = components.get("rewards_penalties", [])
    fee = components.get("management_fee", {})

    claims_requested = sum(_number(x.get("requested_amount")) for x in claims)
    claims_approved = sum(_number(x.get("approved_amount")) for x in claims)
    claims_settlement = sum(_number(x.get("settlement_amount")) for x in claims)
    reward_amount = sum(_number(x.get("amount")) for x in rewards_penalties if x.get("type") == "奖励")
    penalty_amount = sum(_number(x.get("amount")) for x in rewards_penalties if x.get("type") == "罚款")
    penalty_deducted = sum(_number(x.get("deducted_from_settlement")) for x in rewards_penalties if x.get("type") == "罚款")
    contract_fee = _number(fee.get("contract_base")) * _number(fee.get("contract_rate"))
    settlement_fee = _number(fee.get("settlement_base")) * _number(fee.get("settlement_rate"))
    approved_fee = _number(fee.get("settlement_base")) * _number(fee.get("approved_rate", fee.get("contract_rate")))

    return {
        "contract_total": _money(_number(financials.get("contract_total"))),
        "approved_change_total": _money(_number(financials.get("approved_change_total"))),
        "claims_requested": _money(claims_requested),
        "claims_approved": _money(claims_approved),
        "claims_settlement": _money(claims_settlement),
        "unapproved_claim_amount": _money(max(0, claims_settlement - claims_approved)),
        "reward_amount": _money(reward_amount),
        "penalty_amount": _money(penalty_amount),
        "penalty_deducted": _money(penalty_deducted),
        "penalty_omission": _money(max(0, penalty_amount - penalty_deducted)),
        "management_fee_contract": _money(contract_fee),
        "management_fee_settlement": _money(settlement_fee),
        "management_fee_approved": _money(approved_fee),
        "management_fee_excess": _money(max(0, settlement_fee - approved_fee)),
        "settlement_total": _money(_number(financials.get("settlement_total"))),
    }


def detect_cost_component_issues(data: dict[str, Any]) -> list[dict[str, Any]]:
    """Detect unsupported claims, reward/penalty omissions and fee drift."""
    project = data.get("project", {})
    components = project.get("cost_components", {})
    issues: list[dict[str, Any]] = []

    for claim in components.get("claims", []):
        settlement_amount = _number(claim.get("settlement_amount"))
        approved_amount = _number(claim.get("approved_amount"))
        unsupported = max(0, settlement_amount - approved_amount)
        if unsupported <= 0:
            continue
        reference = claim.get("reference_change_id")
        relation = f"，且与变更 {reference} 存在费用重叠" if reference else ""
        issues.append(_issue(
            f"post-claim-{claim.get('id', 'unknown')}",
            "事后",
            "索赔",
            claim.get("title", "索赔金额异常"),
            "high" if reference else "medium",
            unsupported,
            f"结算索赔金额 {settlement_amount:,.0f} 元 > 已批准金额 {approved_amount:,.0f} 元{relation}",
            f"索赔已进入结算，但批准金额不足，未批准部分为 {unsupported:,.0f} 元{relation}。",
            "补充责任划分、工期影响、费用测算和正式审批；未批准部分不得直接计入最终结算。",
            _component_item(claim, claim.get("id", "CLAIM"), "索赔资料"),
        ))

    for item in components.get("rewards_penalties", []):
        amount = _number(item.get("amount"))
        if item.get("type") == "奖励" and not item.get("approved", False) and amount > 0:
            issues.append(_issue(
                f"post-reward-{item.get('id', 'unknown')}",
                "事后",
                "奖罚款",
                item.get("title", "奖励款审批缺失"),
                "medium",
                amount,
                "奖励款 approved=false 但已进入结算",
                "结算书计入奖励款，但未见合同约定奖励确认单或审批记录。",
                "补充奖励条件达成证明和审批单；证据不足部分不予计入。",
                _component_item(item, item.get("id", "REWARD"), "奖罚款资料"),
            ))
        if item.get("type") == "罚款":
            deducted = _number(item.get("deducted_from_settlement"))
            omission = max(0, amount - deducted)
            if item.get("approved", False) and omission > 0:
                issues.append(_issue(
                    f"post-penalty-{item.get('id', 'unknown')}",
                    "事后",
                    "奖罚款",
                    item.get("title", "已确认罚款未扣减"),
                    "high",
                    omission,
                    f"已批准罚款 {amount:,.0f} 元，结算仅扣减 {deducted:,.0f} 元",
                    "已生效的质量处罚未在结算中足额扣减，存在应扣未扣风险。",
                    "将处罚单与结算支付、奖罚台账勾稽，补扣未扣金额。",
                    _component_item(item, item.get("id", "PENALTY"), "奖罚款资料"),
                ))

    fee = components.get("management_fee", {})
    settlement_rate = _number(fee.get("settlement_rate"))
    approved_rate = _number(fee.get("approved_rate", fee.get("contract_rate")))
    excess = max(0, _number(fee.get("settlement_base")) * (settlement_rate - approved_rate))
    if excess > 0:
        fee_item = {"evidence": fee.get("evidence", [])}
        issues.append(_issue(
            "post-fee-001",
            "事后",
            "取费",
            "管理费取费费率偏离合同",
            "medium",
            excess,
            f"结算管理费率 {settlement_rate:.2%} > 合同/批准费率 {approved_rate:.2%}",
            "结算按高于合同约定的费率计取管理费，超额部分缺少有效批准依据。",
            "核对合同取费条款、批准费率和计费基数，核减未获批准的超额管理费。",
            _component_item(fee_item, "MANAGEMENT-FEE", "取费资料"),
        ))
    return issues


def score_fraud_risk(data: dict[str, Any], issues: list[dict[str, Any]]) -> dict[str, Any]:
    """Score historical false-settlement patterns with deterministic features."""
    patterns = {x.get("id"): x for x in data.get("historical_patterns", [])}
    hits: list[dict[str, Any]] = []
    titles = " ".join(x.get("title", "") for x in issues)
    categories = {x.get("category") for x in issues}
    checks = [
        ("duplicate_billing", "工程量" in categories and "重复计费" in titles, "发现同编码同部位重复计价疑点"),
        ("unsupported_item", "虚列子目" in titles, "发现结算申报与实际完成量不一致"),
        ("unreasonable_claim", "索赔" in categories, "发现未批准或部分批准索赔进入结算"),
        ("unapproved_change", "变更签证" in categories, "发现审批链未闭合的变更进入结算"),
        ("rate_drift", "单价" in categories or "取费" in categories, "发现单价或取费费率偏离合同"),
    ]
    score = 0
    for pattern_id, matched, evidence in checks:
        if not matched:
            continue
        pattern = patterns.get(pattern_id, {})
        weight = _number(pattern.get("weight"))
        score += int(weight)
        hits.append({"id": pattern_id, "name": pattern.get("name", pattern_id), "weight": int(weight), "evidence": evidence})
    score = min(100, score)
    level = "高风险" if score >= 60 else "中风险" if score >= 30 else "低风险"
    return {
        "score": score,
        "level": level,
        "model": "历史虚假结算特征规则模型",
        "hits": hits,
        "explanation": "；".join(x["evidence"] for x in hits) if hits else "当前资料未命中已配置的历史虚假结算特征。",
    }


def train_false_settlement_model(data: dict[str, Any], issues: list[dict[str, Any]]) -> dict[str, Any]:
    """Summarize a lightweight supervised model from historical false-settlement cases."""
    patterns = {item.get("id"): item for item in data.get("historical_patterns", [])}
    cases = data.get("historical_cases", [])
    positive_cases = [case for case in cases if case.get("label") in {"虚假结算", "疑似虚假结算"} or case.get("is_false")]
    negative_cases = [case for case in cases if case not in positive_cases]
    positive_base = max(1, len(positive_cases))
    learned_features = []
    issue_text = " ".join(f"{item.get('category', '')} {item.get('title', '')}" for item in issues)

    for pattern_id, pattern in patterns.items():
        positive_hits = sum(1 for case in positive_cases if pattern_id in set(case.get("features", [])))
        negative_hits = sum(1 for case in negative_cases if pattern_id in set(case.get("features", [])))
        precision = positive_hits / max(1, positive_hits + negative_hits)
        coverage = positive_hits / positive_base
        base_weight = int(_number(pattern.get("weight")))
        learned_weight = max(5, min(40, int(round(base_weight * (0.72 + coverage * 0.52 + precision * 0.18)))))
        learned_features.append({
            "id": pattern_id,
            "name": pattern.get("name", pattern_id),
            "base_weight": base_weight,
            "learned_weight": learned_weight,
            "positive_hits": positive_hits,
            "negative_hits": negative_hits,
            "confidence": round((coverage * 0.65 + precision * 0.35), 3),
            "description": pattern.get("description", ""),
            "current_project_hit": pattern.get("name", "") in issue_text or pattern_id.split("_")[0] in issue_text,
        })

    learned_features.sort(key=lambda item: (item["confidence"], item["learned_weight"]), reverse=True)
    return {
        "model": "历史虚假结算特征训练模型",
        "algorithm": "规则特征 + 历史样本监督打分",
        "sample_count": len(cases),
        "positive_cases": len(positive_cases),
        "negative_cases": len(negative_cases),
        "threshold": 60,
        "features": learned_features,
        "top_features": [item["name"] for item in learned_features[:3]],
        "training_status": "已训练" if cases else "使用内置规则权重",
        "training_note": "样本用于学习特征权重，审计结论仍由规则证据链和主审复核确认。",
    }


def _default_audit_models() -> list[dict[str, Any]]:
    """Return the 11 core business models plus pre-filter (方案 §2)."""
    return [
        # ── 前置过滤器（模型0）─────────────────────────────────────────
        {"id": "M-PRE-FILTER", "phase": "事前", "model_type": "过滤类",
         "business_end": "全端",
         "name": "前置过滤器：工程主体清洗与状态交叉校验模型",
         "categories": ["数据清洗"],
         "input_documents": ["项目统计表", "建造合同表"],
         "key_fields": ["project_no", "status"],
         "rule": "自动剔除ZZ前缀非施工主体、交叉比对财务与商务项目状态",
         "output": "纯净施工项目基准集", "agents": ["制度知识库机器人"]},
        # ── 维度一：对上业主结算与确权（M1.1–M1.4）─────────────────────────
        {"id": "M1.1", "phase": "事中", "model_type": "识别类",
         "business_end": "对上业主结算与确权端",
         "name": "外报量确权率偏离检测模型",
         "categories": ["确权", "产值"],
         "input_documents": ["结算书", "验收资料", "施工日志"],
         "key_fields": ["确权率", "完成产值", "确权产值"],
         "rule": "确权率<95%判定滞后，>110%且部位未完工判定超前确权",
         "output": "确权偏离判定与应确未确资金缺口",
         "agents": ["商务专家", "造价专家"]},
        {"id": "M1.2", "phase": "事后", "model_type": "问题类",
         "business_end": "对上业主结算与确权端",
         "name": "久竣未结超时锁定模型",
         "categories": ["久竣未结"],
         "input_documents": ["合同", "竣工验收资料", "收款凭证"],
         "key_fields": ["完工日期", "验收日期", "结算日期", "合同额", "已收款"],
         "rule": "完工超90天未验或验收超180天未结，锁定项目并计算资金占压损失",
         "output": "久竣未结锁定清单与资金占压利息损失",
         "agents": ["财务专家", "商务专家"]},
        {"id": "M1.3", "phase": "事中", "model_type": "预警类",
         "business_end": "对上业主结算与确权端",
         "name": "变更签证索赔滞后模型",
         "categories": ["变更签证", "索赔"],
         "input_documents": ["变更签证台账", "索赔资料"],
         "key_fields": ["变更日期", "审批日期", "索赔日期"],
         "rule": "变更超14天未确权或索赔超28天未书面提出判定滞后",
         "output": "变更签证滞后清单与索赔权丧失损失测算",
         "agents": ["商务专家", "法务专家"]},
        {"id": "M1.4", "phase": "事后", "model_type": "问题类",
         "business_end": "反舞弊与审计端",
         "name": "生效判决收入冲减模型",
         "categories": ["诉讼", "收入"],
         "input_documents": ["判决书", "财务账套"],
         "key_fields": ["判决金额", "账面收入"],
         "rule": "账面收入>法院判决金额时判定财务信息失真",
         "output": "强制红字冲销凭证与瞒报通报",
         "agents": ["法务专家", "财务专家"]},
        # ── 维度二：对下分包分供结算（M2.1–M2.5）─────────────────────────
        {"id": "M2.1", "phase": "事后", "model_type": "问题类",
         "business_end": "对下分包结算端",
         "name": "分包超结超付与结算超合同5%审批穿透模型",
         "categories": ["分包结算", "审批合规"],
         "input_documents": ["合同", "结算书", "红头纪要"],
         "key_fields": ["合同额", "审定额", "超额率", "三重一大批复"],
         "rule": "超额>5%且缺三重一大批复判定重大违规",
         "output": "违规越权结算判定与支付冻结指令",
         "agents": ["商务专家", "法务专家"]},
        {"id": "M2.2", "phase": "事中", "model_type": "识别类",
         "business_end": "反舞弊与审计端",
         "name": "合同明令禁止签证项与清单外虚假结算穿透模型",
         "categories": ["虚假结算", "签证"],
         "input_documents": ["合同", "结算书", "业主清单"],
         "key_fields": ["禁止签证标志", "结算项", "清单外项"],
         "rule": "包干合同出现签证或清单外无对应子目判定虚假结算",
         "output": "违规签证与虚假列项清单及应扣减金额",
         "agents": ["虚假结算识别智能体", "造价专家"]},
        {"id": "M2.3", "phase": "事后", "model_type": "问题类",
         "business_end": "对下分包结算端",
         "name": "材料超耗未扣硬算模型",
         "categories": ["跑冒滴漏"],
         "input_documents": ["材料台账", "结算书", "合同"],
         "key_fields": ["预算量", "实际消耗", "超耗率", "应扣金额"],
         "rule": "钢筋超耗>2%或混凝土超耗>0%时按150%单价硬算应扣金额",
         "output": "超耗扣减差额与强制补扣指令",
         "agents": ["造价专家", "物资专家"]},
        {"id": "M2.4", "phase": "事中", "model_type": "预警类",
         "business_end": "招采供应链端",
         "name": "未签合同先进场施工时序倒置检测模型",
         "categories": ["时序倒置"],
         "input_documents": ["合同", "施工日志"],
         "key_fields": ["合同签订日", "施工日志首日"],
         "rule": "合同签订日晚于施工日志首日判定时序倒置",
         "output": "违规进场判定与责任锁定",
         "agents": ["招采供应链专家", "履约专家"]},
        {"id": "M2.5", "phase": "事后", "model_type": "问题类",
         "business_end": "对下分包结算端",
         "name": "分包甩项代工与负结算未结清收模型",
         "categories": ["代工", "负结算"],
         "input_documents": ["代工台账", "结算书"],
         "key_fields": ["代工金额", "审定额"],
         "rule": "代工费用未100%扣减或审定额<0时判定债权风险",
         "output": "代工漏扣补扣与负结算全局冻结指令",
         "agents": ["商务专家", "财务专家"]},
        # ── 维度三：业财法工全域穿透（M3.1–M3.3）─────────────────────────
        {"id": "M3.1", "phase": "事后", "model_type": "问题类",
         "business_end": "财务审计端",
         "name": "存货未报耗虚增利润模型",
         "categories": ["存货", "利润"],
         "input_documents": ["SAP总账", "竣工验收资料"],
         "key_fields": ["存货余额", "验收日期"],
         "rule": "竣工超30天原材料余额>0判定虚增利润",
         "output": "虚增利润金额与强制调整通知",
         "agents": ["财务专家", "制度知识库机器人"]},
        {"id": "M3.2", "phase": "事后", "model_type": "识别类",
         "business_end": "财务审计端",
         "name": "隐性贴息利息侵蚀与全口径真实效益还原模型",
         "categories": ["贴息", "真实效益"],
         "input_documents": ["财务总账", "利润表", "贴息台账"],
         "key_fields": ["账面利润", "贴息", "逾期利息", "借款利息"],
         "rule": "账面盈利但真实利润<0判定虚盈实亏",
         "output": "全口径真实效益还原报告与责任成本考核调整",
         "agents": ["财务专家", "主审驾驶舱"]},
        {"id": "M3.3", "phase": "事中", "model_type": "预警类",
         "business_end": "法务风控端",
         "name": "招采付款条件倒挂与拖欠高压线强制止损模型",
         "categories": ["付款倒挂", "拖欠"],
         "input_documents": ["招采合同", "业主付款记录"],
         "key_fields": ["对下比例", "对上比例", "拖欠比", "拖欠月数"],
         "rule": "对下>对上判定倒挂；拖欠>30%且>12月触发强制止损",
         "output": "招采倒挂预警与强制停工熔断指令",
         "agents": ["法务专家", "主审驾驶舱"]},
    ]

def run_core_models(data: dict[str, Any]) -> dict[str, list[dict[str, Any]]]:
    """Execute all 11 core business models + pre-filter on project data."""
    import datetime
    project = data.get("project", {}) or {}
    financials = project.get("financials", {}) or {}
    results: dict[str, list[dict[str, Any]]] = {m["id"]: [] for m in _default_audit_models()}

    def _parse_date(d: Any) -> datetime.date | None:
        if not d: return None
        try:
            if isinstance(d, datetime.date): return d
            return datetime.datetime.strptime(str(d)[:10], "%Y-%m-%d").date()
        except: return None
    today = datetime.date.today()

    # Pre-filter
    p_id, p_st = str(project.get("id","")), str(project.get("status",""))
    if p_id.startswith("ZZ") or p_st in {"关闭","已注销","冻结"}:
        results["M-PRE-FILTER"].append({"hit":True,"title":"非施工主体或状态异常","amount":0,
            "detail":f"项目{p_id}含ZZ前缀或状态={p_st},已过滤","risk":"高风险"})

    # M1.1
    compl = _number(project.get("completed_output"))
    conf = _number(project.get("confirmed_output"))
    if compl>0 and conf>0:
        pct = conf/compl
        if pct<0.95:
            gap=_money(compl-conf)
            results["M1.1"].append({"hit":True,"title":"确权率偏低","amount":gap,
                "detail":f"确权率{pct:.1%}<95%,应确未确缺口约{gap:,}元","risk":"中风险"})
        elif pct>1.10:
            results["M1.1"].append({"hit":True,"title":"超前确权","amount":_money(conf-compl),
                "detail":f"确权率{pct:.1%}>110%,可能存在超前确权","risk":"中风险"})

    # M1.2
    finish=_parse_date(project.get("actual_finish_date"))
    accept=_parse_date(project.get("acceptance_date"))
    settle=_parse_date(project.get("settlement_date"))
    if finish and (not accept or (accept-finish).days>90):
        contract_amt=_number(financials.get("contract_total"))
        pd=(today-finish).days if not accept else (accept-finish).days
        loss=_money(contract_amt*0.05*min(pd,365)/365) if pd>90 else 0
        results["M1.2"].append({"hit":True,"title":"完工超90天未验收","amount":loss or _money(contract_amt*0.05),
            "detail":f"完工{finish},验收{accept or '无'},已逾{pd}天,资金占压损失约{loss:,}元","risk":"高风险"})
    if accept and (not settle or (today-accept).days>180):
        c2=_number(financials.get("contract_total"))
        coll=_number(project.get("cumulative_collection"))
        results["M1.2"].append({"hit":True,"title":"验收超180天未结算","amount":_money(c2-coll),
            "detail":f"验收{accept}已逾{(today-accept).days}天,未收款余额{_money(c2-coll):,}元","risk":"高风险"})

    # M1.3
    for var in project.get("variations",[]):
        vd=_parse_date(var.get("occur_date"))
        if vd and (today-vd).days>14 and var.get("status")!="APPROVED":
            results["M1.3"].append({"hit":True,"title":f"变更{var.get('id','')}确权滞后",
                "amount":_money(_number(var.get("amount"))),
                "detail":f"变更{var.get('id','')}已发生{(today-vd).days}天,审批状态={var.get('status')}","risk":"中风险"})
    if not results["M1.3"]:
        for cl in project.get("cost_components",{}).get("claims",[]):
            if cl.get("status") in {"待审批","部分批准"} and _number(cl.get("requested_amount"))>0:
                results["M1.3"].append({"hit":True,"title":f"索赔{cl.get('id','')}审批滞后",
                    "amount":_money(_number(cl.get("requested_amount"))),
                    "detail":f"索赔{cl.get('id','')}申报{_number(cl.get('requested_amount')):,.0f}元,状态={cl.get('status')}","risk":"高风险"})

    # M1.4
    judgment=_number(project.get("court_judgment_amount"))
    book_rev=_number(project.get("sap_cumulative_revenue"))
    if judgment>0 and book_rev>judgment:
        results["M1.4"].append({"hit":True,"title":"判决收入未冲减","amount":_money(book_rev-judgment),
            "detail":f"判决{judgment:,.0f}元,账面{book_rev:,.0f}元,虚增{book_rev-judgment:,.0f}元","risk":"高风险"})

    # M2.1
    sub_ct=_number(project.get("subcontract_settlement",{}).get("contract_amount"))
    sub_au=_number(project.get("subcontract_settlement",{}).get("audit_amount"))
    if sub_ct>0 and sub_au>0 and sub_au>sub_ct*1.05:
        triple=(project.get("approval_minutes",{}) or {}).get("is_triple_importance_large",False)
        results["M2.1"].append({"hit":True,"title":"结算超合同5%且缺三重一大批复",
            "amount":_money(sub_au-sub_ct),
            "detail":f"分包结算{sub_au:,.0f}vs合同{sub_ct:,.0f},超额{sub_au/sub_ct-1:.1%},三重一大={triple}","risk":"高风险"})

    # M2.2
    if project.get("no_variation_allowed",False):
        for ln in data.get("line_items",[]):
            if _number(ln.get("change",{}).get("amount"))>0:
                results["M2.2"].append({"hit":True,"title":"包干合同出现签证项",
                    "amount":_money(_number(ln.get("change",{}).get("amount"))),
                    "detail":f"包干合同(禁止签证),清单{ln.get('item_code','')}存在变更","risk":"高风险"})

    # M2.3
    for mat in project.get("materials",[]):
        bq=_number(mat.get("budget_quantity")); aq=_number(mat.get("actual_quantity"))
        pr=_number(mat.get("purchase_unit_price")); nm=mat.get("name","")
        if bq>0 and aq>bq:
            rate=(aq-bq)/bq; lim=0.02 if "钢" in nm else 0.00
            if rate>lim:
                deduct=_money((aq-bq)*pr*1.5)
                results["M2.3"].append({"hit":True,"title":f"{nm}超耗未扣","amount":deduct,
                    "detail":f"{nm}预算{bq:,.0f}实际{aq:,.0f}超耗{rate:.1%},按150%应扣{deduct:,}元","risk":"中风险"})
# M2.4
    sign=_parse_date(project.get("sign_date"))
    diary=_parse_date(project.get("diary_first_date"))
    if sign and diary and diary<sign:
        results["M2.4"].append({"hit":True,"title":"未签合同先进场","amount":0,
            "detail":f"合同签订日{sign}晚于施工日志首日{diary},时序倒置{(sign-diary).days}天","risk":"中风险"})

    # M2.5
    for lb in project.get("third_party_labors",[]):
        la=_number(lb.get("amount")); ded=lb.get("is_deducted_in_settlement",False)
        if la>0 and not ded:
            results["M2.5"].append({"hit":True,"title":f"代工{lb.get('name','')}未扣减","amount":_money(la),
                "detail":f"{lb.get('name','')}代工{la:,.0f}元,结算未扣减","risk":"高风险"})
    sub_au2=_number(project.get("subcontract_settlement",{}).get("audit_amount"))
    if sub_au2<0:
        results["M2.5"].append({"hit":True,"title":"负结算未清收","amount":_money(abs(sub_au2)),
            "detail":f"分包审定{sub_au2:,.0f}元,债权风险","risk":"高风险"})

    # M3.1
    invb=_number(project.get("sap_inventory_balance"))
    if accept and invb>0 and (today-accept).days>30:
        results["M3.1"].append({"hit":True,"title":"竣工后存货未报耗","amount":_money(invb),
            "detail":f"验收已{(today-accept).days}天,存货余额{invb:,.0f}元,虚增利润","risk":"高风险"})

    # M3.2
    bp=_number(project.get("book_profit"))
    disc=_number(project.get("supply_chain_discount_fee"))
    oi=_number(project.get("steel_overdue_interest"))
    li=_number(project.get("internal_loan_interest"))
    tp=bp-disc-oi-li
    if bp>0 and tp<0:
        results["M3.2"].append({"hit":True,"title":"虚盈实亏","amount":_money(abs(tp)),
            "detail":f"账面{bp:,.0f},扣贴息{disc:,.0f}+逾期{oi:,.0f}+借款{li:,.0f},真实{tp:,.0f}","risk":"高风险"})
    elif disc>0 or oi>0 or li>0:
        ero=_money(disc+oi+li)
        results["M3.2"].append({"hit":True,"title":"隐性贴息侵蚀效益","amount":ero,
            "detail":f"账面{bp:,.0f},隐性侵蚀{ero:,}元(贴息{disc:,.0f}+逾期{oi:,.0f}+借款{li:,.0f})","risk":"中风险"})

    # M3.3
    up=_number(project.get("owner_progress_pay_ratio"))
    dn=_number(project.get("subcontract_pay_ratio"))
    dr=_number(project.get("owner_default_ratio"))
    dm=int(_number(project.get("owner_default_months")))
    if dn>up:
        results["M3.3"].append({"hit":True,"title":"付款条件倒挂","amount":0,
            "detail":f"对下{dn:.0%}>对上{up:.0%},资金链紧张","risk":"中风险"})
    if dr>0.30 and dm>12:
        results["M3.3"].append({"hit":True,"title":"业主拖欠超高压线",
            "amount":_money(dr*_number(financials.get("contract_total"))),
            "detail":f"业主拖欠{dr:.0%}已{dm}月,触发强制止损","risk":"高风险"})

    return results
def build_model_chain(core_results: dict[str, list[dict[str, Any]]],
                      issues: list[dict[str, Any]] | None = None) -> dict[str, Any]:
    """Build the 8-step execution pipeline (模型链) grouping models by execution order.
    方案 §3.1 '双轨为经、穿透为纬'三维多模型关联图谱."""
    issues = issues or []
    chain_steps = [
        {"step": 0, "name": "前置过滤与数据清洗",
         "models": ["M-PRE-FILTER"], "model_names": ["前置过滤器"],
         "description": "剔除ZZ实体，修复项目状态漂移，输出纯净施工项目基准"},
        {"step": 1, "name": "准入合规性扫描",
         "models": ["M2.4", "M2.2"],
         "model_names": ["未签合同先进场检测", "禁止签证与清单外虚假结算穿透"],
         "description": "扫描未签先进场、禁止签证项与清单外虚假列项"},
        {"step": 2, "name": "材料超耗与代工漏扣硬算",
         "models": ["M2.3", "M2.5"],
         "model_names": ["材料超耗未扣硬算", "分包代工负结算清收"],
         "description": "确定性硬算材料超耗150%倍扣减、分包甩项代工漏扣与负结算"},
        {"step": 3, "name": "结算合规性审查",
         "models": ["M2.1"],
         "model_names": ["结算超5%审批穿透"],
         "description": "审查最终结算合规性，超额5%三重一大穿透"},
        {"step": 4, "name": "对上确权与久竣未结审查",
         "models": ["M1.1", "M1.2", "M1.3"],
         "model_names": ["确权率检测", "久竣未结锁定", "变更签证滞后"],
         "description": "确权率<95%、久竣未结超时、变更签证滞后审查"},
        {"step": 5, "name": "业财法工穿透审计",
         "models": ["M3.1", "M3.2", "M1.4"],
         "model_names": ["存货虚增利润", "隐性贴息还原", "判决收入冲减"],
         "description": "存货账实对账、隐性贴息侵蚀、生效判决收入冲销"},
        {"step": 6, "name": "九大管理归因与战略定论链串联",
         "models": ["G1", "G2", "G3", "G4", "G5", "G6", "G7", "G8", "G9"],
         "model_names": ["先干后谈责任错配", "结算滞后虚盈实亏", "证据链断点阻断",
                         "层层转包价格失控", "盲目垫资止损击穿", "三重一大审批失守",
                         "诉讼判决未冲销", "贴息反噬效益侵蚀", "时空穿越人证分离"],
         "description": "将11个微观模型升维聚合为九大高管视角管理归因与战略审计定论链（商务结算合规4链+公司治理与战略风控5链），输出三段式公文定论"},
        {"step": 7, "name": "合规监管映射与处置建议",
         "models": [], "model_names": [],
         "description": "自动匹配违规行为分类与一/二/三级处置建议"},
        {"step": 8, "name": "主审驾驶舱输出",
         "models": [], "model_names": [],
         "description": "自动生成结算审计报告+整改销号工单包"},
    ]
    total_hits = 0
    for step in chain_steps:
        hits = []
        hit_amount = 0
        for mid in step["models"]:
            findings = core_results.get(mid, [])
            for f in findings:
                hits.append(f)
                hit_amount += _money(_number(f.get("amount")))
        step["hit_count"] = len(hits)
        step["hit_amount"] = hit_amount
        step["hits"] = hits[:5]
        total_hits += len(hits)
    return {
        "steps": chain_steps,
        "total_hits": total_hits,
        "pipeline_status": "正常流转" if total_hits > 0 else "无命中",
    }
def build_nine_grid(core_results: dict[str, list[dict[str, Any]]],
                    data: dict[str, Any]) -> dict[str, Any]:
    """构建九宫格（责任/风险/处置等级判定）. 方案 §3.5 商务结算"风险-效益"九宫格离散决策矩阵."""
    project = data.get("project", {}) or {}
    financials = project.get("financials", {}) or {}

    # 结算风险指数 R
    risk_weights = {"M2.1":3.0,"M2.5":2.5,"M1.4":2.5,"M2.2":2.0,"M1.2":2.0,
                    "M3.1":2.0,"M1.1":1.5,"M2.3":1.5,"M2.4":1.0,"M1.3":1.0,"M3.2":1.0,"M3.3":1.0}
    hit_risk = sum(risk_weights.get(mid,1.0) for mid, fs in core_results.items() if fs)
    hit_count = sum(1 for fs in core_results.values() if fs)
    r_score = round(min(max(1.0, hit_risk/hit_count if hit_count>0 else 1.0), 3.0), 1)

    # 真实效益指数 E
    contract = _number(financials.get("contract_total"))
    certified = _number(financials.get("actual_certified_total"))
    e_raw = certified/contract if contract>0 else 0.5
    e_score = round(min(max(1.0, e_raw*3), 3.0), 1)
    if any(h.get("title")=="虚盈实亏" for h in core_results.get("M3.2",[])):
        e_score = max(1.0, e_score-0.5)

    r_idx = min(int(r_score-1), 2)
    e_idx = min(int(3-e_score), 2)  # high E = low row (0=best)
    quadrant = f"({r_idx+1},{e_idx+1})"
    zones = {
        (0,0):("双优区","优质项目","保持"), (0,1):("观察区","效益好但风险上升","关注"),
        (0,2):("重点观察区","效益好但高风险","严控"),
        (1,0):("稳定区","风险中等+效益好","维持"), (1,1):("观察区","风险中等+效益中等","核查"),
        (1,2):("危险区","风险高+效益中等","限期整改"),
        (2,0):("重点观察区","风险低但效益差","帮扶"), (2,1):("观察区","确权滞后+微利微亏","限期3月确权"),
        (2,2):("危险区","高风险+低效益","立即止损追责"),
    }
    zone_info = zones.get((r_idx, e_idx), ("未知","",""))

    risk_details = []
    benefit_details = []
    for mid, findings in core_results.items():
        if not findings: continue
        model = next((m for m in _default_audit_models() if m["id"]==mid), None)
        amt = sum(_number(f.get("amount")) for f in findings)
        if any(f.get("risk")=="高风险" for f in findings):
            risk_details.append({"model_id":mid,"name":model["name"] if model else mid,
                                 "amount":_money(amt),"risk":"高风险"})
        if mid in {"M3.2","M3.1","M1.4"}:
            benefit_details.append({"model_id":mid,"name":model["name"] if model else mid,
                                    "amount":_money(abs(amt)),"impact":"效益扣减"})
    color = "red" if r_idx>=2 else "orange" if r_idx>=1 else "yellow" if e_idx>=1 else "green"
    return {
        "r_score":r_score, "e_score":e_score, "quadrant":quadrant,
        "zone":zone_info[0], "diagnosis":zone_info[1], "action":zone_info[2],
        "risk_details":risk_details, "benefit_details":benefit_details,
        "summary":{"risk_level":"高风险" if r_score>=2.5 else "中风险" if r_score>=1.5 else "低风险",
                   "benefit_level":"高效益" if e_score>=2.5 else "中等效益" if e_score>=1.5 else "低效益",
                   "zone_color":color},
    }


def build_cross_model_hints(core_results: dict[str, list[dict[str, Any]]],
                            issues: list[dict[str, Any]] | None = None) -> dict[str, Any]:
    """构建九大管理归因与战略审计定论链（G1~G9）及命中聚合.
    方案 §3.2 九大管理归因与战略审计定论链：彻底摒弃三条关联提示链，
    将 11 个微观计算模型升维聚合成商务结算合规（4 链）与公司治理与战略风控（5 链）两大阵列."""
    issues = issues or []
    chains = [
        {"chain_id":"G1","section":"商务结算合规",
         "name":"先干后谈与未定价变更责任错配链",
         "path":"口头/会议越权指令 → 模型2.4(未签合同先进场) → 模型1.3(变更超14天未办经济确权) → 模型2.2(结算申报清单外项被全额核减)",
         "description":"现场“必须马上干”与商务“以后审价”两次决定脱节：工程端取即时进度、商务端保留审价空间，现场材料已采购、隐蔽已覆盖，证据灭失、成本混淆入原合同，项目部被迫承担现金垫资与结算让价双重损失。",
         "models":["M2.4","M1.3","M2.2"],
         "model_names":["未签合同先进场","变更超14天未办经济确权","清单外项被全额核减"],
         "conclusion":{
             "qualitative":"项目违规执行“先干后谈”，存在严重的时序错配与管理责任转嫁；",
             "quantitative":"未定价变更先行施工2.4个月，导致209万元新增成本混入原合同且缺乏有效图纸和签认证据，结算时被迫折让核减180万元；",
             "management":"暴露出上级越权下达口头指令而回避经济后果的坏管理，系统依据“一页最低确认机制”将指令发起人锁定为风险第一责任人。"},
         "severity":"重大","action":"锁定指令发起人为第一责任人并追责"},
        {"chain_id":"G2","section":"商务结算合规",
         "name":"结算滞后与虚盈实亏时间亏损链",
         "path":"模型1.1(外报量确权偏离) → 模型1.2(久竣未结超期515天) → 模型3.1(存货未报耗挂账868万) → 模型3.2(产生贴息利息8424万)",
         "description":"为冲产值提前确认虚高收入，竣工后资料不齐、签证未闭合拖延结算，形成大额“应收未结”挂账，账面利润被年化资金利息(3.5%~5%)吞噬，形成典型“时间型亏损”。",
         "models":["M1.1","M1.2","M3.1","M3.2"],
         "model_names":["外报量确权偏离","久竣未结超期锁定","存货未报耗虚增利润","隐性贴息侵蚀还原"],
         "conclusion":{
             "qualitative":"项目表面稳健盈利，实质已陷入流动性枯竭的“虚盈实亏”经营陷阱；",
             "quantitative":"竣工超期未结导致2.1亿元资金沉淀，产生贴息利息8424万元，叠加868万元原材料未报耗虚增利润，真实效益由账面盈利1200万沦为实质巨亏-3022万元；",
             "management":"暴露出“以产值为导向”的考核扭曲，必须强制将“资金占用成本与折现率”纳入利润核算，推动结算全面提速。"},
         "severity":"重大","action":"强制资金占用成本与折现率纳入利润核算"},
        {"chain_id":"G3","section":"商务结算合规",
         "name":"证据治理与六大证据链断点阻断链",
         "path":"六链拓扑扫描(合同链/变更链/签证链/工程量链/费用链/工期链) → 证据计数evidence_count<2 → 模型2.2(禁止签证项违规) → 模型1.4(法院因缺证调减)",
         "description":"结算争议的本质是“证据贫困”：签证缺监理签字、隐蔽照片无GPS水印、合同审批意见未落入正文，真实施工的事项在法律与结算中无法被采信。",
         "models":["M2.2","M1.4"],
         "model_names":["禁止签证项违规","法院因缺证调减"],
         "conclusion":{
             "qualitative":"项目过程证据治理严重缺失，核心权利主张缺乏法理证据支撑；",
             "quantitative":"六大证据链共检出15份签证缺三方签认、3处隐蔽工程照片缺失时空水印，导致590万元变更款在结算中被业主刚性剔除；",
             "management":"必须将审计关口前移至中期审计，建立“证据质量指数(EQI)”考核与预防性备案机制，确保无证据事项在系统内无法流转。"},
         "severity":"重大","action":"关口前移中期审计+EQI证据质量考核"},
        {"chain_id":"G4","section":"商务结算合规",
         "name":"层层转包与影子分包价格失控链",
         "path":"模型3.3(招采条件倒挂) → 模型2.1(分包结算超5%缺审批) → 模型2.3(材料超耗150%漏扣) → 模型2.5(第三方代工733万未扣)",
         "description":"分包准入失真(个人施工队挂靠资质)、合同条款松散、同类工序班组间价差20%~35%随意认价、过程管理责任断裂，利润在多层转包中被彻底稀释。",
         "models":["M3.3","M2.1","M2.3","M2.5"],
         "model_names":["招采付款条件倒挂","分包结算超5%缺审批","材料超耗150%漏扣","第三方代工未扣"],
         "conclusion":{
             "qualitative":"分包管理存在严重的资质挂靠、随意认价与跑冒滴漏失控；",
             "quantitative":"人工费单价偏离区域限价25%，第三方代工733万元与材料超耗款未在结算中扣减，导致单项目分包超结超付3160万元；",
             "management":"暴露出项目经理权力过度集中与分包限价库监管缺位，建立分包商黑名单，并对失职商务人员予以二级经济赔偿追责。"},
         "severity":"重大","action":"建立分包商黑名单+二级经济赔偿追责"},
        {"chain_id":"G5","section":"公司治理与战略风控",
         "name":"盲目垫资与履约惯性击穿强制止损红线链",
         "path":"模型3.3(招采条件倒挂) → 业主逾期拖欠>30%且>12个月 → 项目持续采购施工 → 产生巨额负现金流",
         "description":"受“保产值、冲营收”考核驱使，业主已实质违约甚至暴雷仍产生盲目履约惯性，不敢停工止损，持续对下招采垫资，企业沦为高风险业主的融资工具。",
         "models":["M3.3"],
         "model_names":["招采倒挂与强制止损"],
         "conclusion":{
             "qualitative":"违背审慎经营原则，击穿营销风控底线与停缓建强制止损红线；",
             "quantitative":"在业主拖欠2.48亿元且逾期超14个月的情况下，项目部未按规定启动强制停工，反向垫资6314万元；",
             "management":"暴露出短期产值考核对风险防控的严重扭曲，建议对分管领导及项目经理追究违规经营投资责任。"},
         "severity":"重大","action":"追究违规经营投资责任"},
        {"chain_id":"G6","section":"公司治理与战略风控",
         "name":"“三重一大”与分级授权在结算审批中的制度性失守链",
         "path":"模型2.1(分包结算超合同额5%) → approval_minutes穿透 → 发现无党委会/董事会纪要 → 演变为“事后补签或班子会代行”",
         "description":"超额数十倍的分包结算款流出本应由党委会/董事会作为“三重一大”前置集体审议，实际被项目经理或个别领导内部拍板替代，分级授权手册(DOA)被架空。",
         "models":["M2.1"],
         "model_names":["分包结算超5%审批穿透"],
         "conclusion":{
             "qualitative":"大额分包结算审批违反“三重一大”法定程序，内部控制实质性失效；",
             "quantitative":"安装工程等11个分包结算超合同总额3160万元(超额率达51.8%)，均未履行“三重一大”集体决议程序；",
             "management":"暴露出权力失衡与合规监督虚化，已形成重大资金流失漏洞，对越权签字人员移交纪检监察部门追责。"},
         "severity":"重大","action":"移交纪检监察部门追责"},
        {"chain_id":"G7","section":"公司治理与战略风控",
         "name":"诉讼判决瞒报与财务收入虚增失真链",
         "path":"模型1.4(生效判决与结算审减收入冲减) → 法院生效判决调减结算确权 → 项目部瞒报 → 财务未冲减虚增收入 → 司法强划暴雷",
         "description":"商务与法务收到不利诉讼判决后选择性隐瞒，财务账面仍按原高估确权挂账虚增报表利润，直至企业银行账户被法院突然划扣。",
         "models":["M1.4"],
         "model_names":["生效判决收入冲减"],
         "conclusion":{
             "qualitative":"存在隐瞒法律败诉风险、财务会计信息严重失真的恶性违规行为；",
             "quantitative":"法院生效判决调减结算确权748.68万元，项目部隐瞒未报，财务未作收入冲减，导致账面持续虚增利润超18个月；",
             "management":"暴露出部门壁垒下的“报喜不报忧”风气，严重误导管理层经营决策，责令立即冲销账面虚增收入并通报追责。"},
         "severity":"重大","action":"责令冲销虚增收入并通报追责"},
        {"chain_id":"G8","section":"公司治理与战略风控",
         "name":"供应链贴息反噬与全口径真实效益侵蚀链",
         "path":"大量开具供应链金融凭证(云信/融易达) → 巨额贴息与逾期利息 → 模型3.2(全口径效益还原) → 财务总账侵蚀 → 真实巨亏",
         "description":"为缓解资金压力过度使用供应链金融工具，贴息与利息被计入财务费用总账、未摊入项目制造成本，掩盖项目真实亏损。",
         "models":["M3.2"],
         "model_names":["隐性贴息侵蚀与真实效益还原"],
         "conclusion":{
             "qualitative":"隐性资金成本严重反噬经营成果，项目呈现典型的“账面虚盈、实质巨亏”；",
             "quantitative":"账面显示盈利1200万元，但产生供应链贴息3713万元与钢材逾期利息4710万元(合计侵蚀8424万元)，全口径真实效益为实质亏损-3022万元；",
             "management":"成本核算口径不完整，必须建立“全口径成本核算机制”，将贴息利息全额计入项目责任成本。"},
         "severity":"重大","action":"建立全口径成本核算机制"},
        {"chain_id":"G9","section":"公司治理与战略风控",
         "name":"“时空穿越”人证分离与现场安全实质悬空链",
         "path":"视频时空碰撞模型(时间/空间多维冲突) → 同一工作负责人同期在多地工单出现 → 模型2.4(未签先进场) → 资质挂靠借壳",
         "description":"施工单位借资质投标，中标后派出的全是不具备资质的包工头，关键技术管理人员人证分离，现场安全与技术管理制度完全空转。",
         "models":["M2.4"],
         "model_names":["未签合同先进场/借壳进场"],
         "conclusion":{
             "qualitative":"现场关键管理人员履约严重缺位，构成实质性挂靠借壳与安全管理空转；",
             "quantitative":"系统识别出47项工单存在同一负责人同期异地多重作业(时空穿越)，涉及违规施工金额855万元；",
             "management":"现场安全责任体系实质悬空，对施工单位扣减信用分并列入准入黑名单，按合同约定全额扣收违约金。"},
         "severity":"重大","action":"扣信用分、列准入黑名单、全额扣收违约金"},
    ]
    sections = [
        {"section":"商务结算合规","count":4,"chain_ids":["G1","G2","G3","G4"]},
        {"section":"公司治理与战略风控","count":5,"chain_ids":["G5","G6","G7","G8","G9"]},
    ]
    for chain in chains:
        hits = []
        for mid in chain["models"]:
            for f in core_results.get(mid, []):
                hits.append({"model_id":mid,"title":f.get("title",""),"risk":f.get("risk",""),
                             "amount":_money(_number(f.get("amount")))})
        chain["hit_models"] = list({h["model_id"] for h in hits})
        chain["hit_count"] = len(hits)
        chain["all_hit"] = len(chain["hit_models"]) >= len(chain["models"])
        chain["hit_amount"] = sum(h["amount"] for h in hits)
        chain["hits"] = hits
    matched = [c["chain_id"] for c in chains if c["hit_count"] > 0]
    return {
        "overview":"系统全面摒弃了原本简单的三条关联提示链，将11个微观计算模型升维聚合成九大高管视角管理归因与战略审计定论链（商务结算合规4链 + 公司治理与战略风控5链）。",
        "sections": sections,
        "chains": chains,
        "total_chains": len(chains),
        "matched_chains": matched,
        "unmatched_chains": [c["chain_id"] for c in chains if c["hit_count"] == 0],
        "pipeline_status": "九大定论链正常流转" if matched else "无命中",
    }

def build_audit_model_catalog(
    data: dict[str, Any],
    issues: list[dict[str, Any]],
    remediation_tasks: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Build a lifecycle model catalog spanning warning, issue, detection and close-out models."""
    configured = data.get("audit_models") or _default_audit_models()
    lifecycle = data.get("lifecycle_findings", [])
    remediation_summary = (remediation_tasks or {}).get("summary", {})
    models = []
    for model in configured:
        categories = set(model.get("categories", []))
        matched_issues = [
            issue for issue in issues
            if issue.get("category") in categories or (model.get("id") in {"M-POST-004", "M2.4"} and issue.get("risk") == "高风险")
        ]
        matched_lifecycle = [
            item for item in lifecycle
            if item.get("phase") == model.get("phase")
            and (item.get("model_type") == model.get("model_type") or item.get("domain") in model.get("business_end", ""))
        ]
        if model.get("model_type") == "整改类":
            hit_count = remediation_summary.get("total", 0)
            risk_amount = remediation_summary.get("amount", 0)
        else:
            hit_count = len(matched_issues) + len(matched_lifecycle)
            risk_amount = sum(_number(item.get("amount")) for item in matched_issues + matched_lifecycle)
        enriched = deepcopy(model)
        enriched.update({
            "status": "已命中" if hit_count else "待监控",
            "hit_count": hit_count,
            "risk_amount": _money(risk_amount),
            "matched_issue_ids": [item.get("id", "") for item in matched_issues],
            "matched_lifecycle_ids": [item.get("id", "") for item in matched_lifecycle],
        })
        models.append(enriched)

    phase_counts = {
        phase: sum(1 for item in models if item.get("phase") == phase)
        for phase in ["事前", "事中", "事后"]
    }
    type_counts = {
        model_type: sum(1 for item in models if item.get("model_type") == model_type)
        for model_type in sorted({item.get("model_type") for item in models})
    }
    return {
        "models": models,
        "summary": {
            "model_count": len(models),
            "active_model_count": sum(1 for item in models if item.get("hit_count", 0) > 0),
            "phase_counts": phase_counts,
            "type_counts": type_counts,
            "business_end_count": len({item.get("business_end") for item in models}),
            "hit_count": sum(item.get("hit_count", 0) for item in models),
            "risk_amount": sum(item.get("risk_amount", 0) for item in models),
        },
    }


def build_agent_orchestration(data: dict[str, Any], issues: list[dict[str, Any]]) -> dict[str, Any]:
    """Assign audit work to the chief auditor, domain experts and assistant bots."""
    experts = {item.get("name"): item for item in data.get("experts", [])}
    roles = [
        {"id": "chief", "name": "主审驾驶舱", "type": "主审", "responsibility": "统一分派疑点、综合专家意见、完成人工复核并签发报告"},
        {"id": "parser", "name": "资料解析智能体", "type": "处理类", "responsibility": "识别招标、合同、签证、竣工、结算及台账文件并抽取字段"},
        {"id": "modeler", "name": "结构化建模智能体", "type": "处理类", "responsibility": "将多文档映射到统一清单、费用、证据和生命周期模型"},
        {"id": "rule", "name": "规则比对智能体", "type": "问题类", "responsibility": "执行工程量、单价、金额、索赔、奖罚、取费和跑冒滴漏规则"},
        {"id": "fraud", "name": "虚假结算识别智能体", "type": "问题类", "responsibility": "基于历史特征识别重复计费、虚列子目和不合理索赔"},
        {"id": "knowledge", "name": "制度知识库机器人", "type": "问询类", "responsibility": "返回制度依据、适用条件、证据清单和补证建议"},
    ]
    roles.extend(
        {
            "id": f"expert-{expert.get('id', index)}",
            "name": expert.get("name", ""),
            "type": "专家类",
            "responsibility": expert.get("domain", ""),
        }
        for index, expert in enumerate(data.get("experts", []), 1)
    )
    category_map = {
        "工程量": ["造价专家", "商务专家", "履约专家"],
        "单价": ["造价专家", "招采供应链专家", "商务专家"],
        "实际完成量": ["履约专家", "造价专家"],
        "变更签证": ["商务专家", "法务专家", "造价专家"],
        "索赔": ["商务专家", "法务专家", "财务专家"],
        "奖罚款": ["商务专家", "法务专家", "财务专家"],
        "取费": ["造价专家", "财务专家", "商务专家"],
        "跑冒滴漏": ["招采供应链专家", "履约专家", "财务专家"],
    }
    tasks = []
    for issue in issues:
        expert_names = category_map.get(issue.get("category"), ["造价专家", "商务专家"])
        task_agents = [
            {"role": "主审驾驶舱", "action": "复核裁量"},
            {"role": "规则比对智能体", "action": "复算金额和规则命中"},
            {"role": "制度知识库机器人", "action": "检索制度依据和补证清单"},
        ]
        task_agents.extend(
            {"role": name, "action": experts.get(name, {}).get("domain", "专业复核")}
            for name in expert_names
        )
        if issue.get("category") in {"工程量", "实际完成量", "变更签证", "索赔", "奖罚款", "取费"}:
            task_agents.append({"role": "虚假结算识别智能体", "action": "匹配历史虚假结算特征"})
        tasks.append({
            "issue_id": issue.get("id", ""),
            "title": issue.get("title", ""),
            "category": issue.get("category", ""),
            "risk": issue.get("risk", ""),
            "amount": issue.get("amount", 0),
            "agents": task_agents,
            "status": "待主审复核",
        })
    return {
        "workflow": ["资料解析", "结构化建模", "规则比对", "疑点分派", "专家会审", "知识问询", "主审复核", "报告签发", "整改闭环"],
        "roles": roles,
        "tasks": tasks,
        "summary": {
            "role_count": len(roles),
            "task_count": len(tasks),
            "expert_count": len(data.get("experts", [])),
            "knowledge_bot_count": 1,
        },
    }


def summarize(
    data: dict[str, Any],
    issues: list[dict[str, Any]],
    comparisons: list[dict[str, Any]],
    cost_aggregation: dict[str, Any] | None = None,
) -> dict[str, Any]:
    financials = data.get("project", {}).get("financials", {})
    cost_aggregation = cost_aggregation or aggregate_cost_components(data)
    risk_total = sum(x["amount"] for x in issues)
    quantity_flags = sum(1 for x in comparisons if abs(x.get("quantity_deviation_settlement_vs_actual") or 0) > DEFAULT_THRESHOLDS["quantity_deviation_pct"])
    price_flags = sum(1 for x in comparisons if abs(x.get("unit_price_deviation_settlement_vs_contract") or 0) > DEFAULT_THRESHOLDS["unit_price_deviation_pct"])
    return {
        "issue_count": len(issues),
        "high_risk_count": sum(1 for x in issues if x["risk"] == "高风险"),
        "medium_risk_count": sum(1 for x in issues if x["risk"] == "中风险"),
        "risk_amount": risk_total,
        "contract_total": _money(_number(financials.get("contract_total"))),
        "approved_change_total": _money(_number(financials.get("approved_change_total"))),
        "settlement_total": _money(_number(financials.get("settlement_total"))),
        "actual_certified_total": _money(_number(financials.get("actual_certified_total"))),
        "paid_total": _money(_number(financials.get("paid_total"))),
        "payment_gap": _money(max(0, _number(financials.get("paid_total")) - _number(financials.get("actual_certified_total")))),
        "quantity_flags": quantity_flags,
        "unit_price_flags": price_flags,
        "document_count": len(data.get("project", {}).get("documents", [])),
        "category_counts": {
            category: sum(1 for item in issues if item.get("category") == category)
            for category in sorted({item.get("category") for item in issues})
        },
        "cost_aggregation": cost_aggregation,
    }


def build_report(
    data: dict[str, Any],
    issues: list[dict[str, Any]],
    summary: dict[str, Any],
    decisions: dict[str, Any] | None = None,
    fraud_assessment: dict[str, Any] | None = None,
    remediation_tasks: dict[str, Any] | None = None,
) -> dict[str, Any]:
    decisions = decisions or {}
    fraud_assessment = fraud_assessment or {}
    remediation_tasks = remediation_tasks or {"summary": {}}
    confirmed = [x for x in issues if decisions.get(x["id"], {}).get("decision") == "confirm"]
    cost = summary.get("cost_aggregation", {})
    remediation_summary = remediation_tasks.get("summary", {})
    return {
        "status": "正式问题待主审签发" if confirmed else "主审复核草稿",
        "title": f"{data.get('project', {}).get('name', '')}工程结算审计报告（草稿）",
        "confirmed_count": len(confirmed),
        "issue_count": len(issues),
        "risk_amount": summary["risk_amount"],
        "confirmed_amount": sum(x["amount"] for x in confirmed),
        "sections": [
            "一、审计范围与资料接入",
            "二、合同价、变更价、实际完成量与结算金额比对",
            "三、异常规则与跑冒滴漏核查",
            "四、审计问题、依据及责任建议",
            "五、整改验证与销号安排",
        ],
        "management_notes": [
            f"识别 {summary['issue_count']} 项疑点，规则测算影响金额 {_money(summary['risk_amount']):,} 元。",
            f"已支付金额较实际完成量对应金额高 {_money(summary['payment_gap']):,} 元，需联动支付审核。",
            f"索赔申报 {cost.get('claims_settlement', 0):,} 元，已批准 {cost.get('claims_approved', 0):,} 元，未批准进入结算 {cost.get('unapproved_claim_amount', 0):,} 元。",
            f"虚假结算特征模型评分 {fraud_assessment.get('score', 0)} 分，等级：{fraud_assessment.get('level', '低风险')}。",
            f"已生成整改任务 {remediation_summary.get('total', 0)} 项，待整改金额 {remediation_summary.get('open_amount', 0):,} 元。",
            "正式审计问题须在主审完成人工裁量后签发。",
        ],
    }


def _remediation_owner(category: str) -> str:
    owner_map = {
        "工程量": "项目商务部 / 造价负责人",
        "单价": "项目商务部 / 招采供应链",
        "实际完成量": "项目履约部 / 工程管理部",
        "变更签证": "项目商务部 / 合同法务",
        "索赔": "项目商务部 / 合同法务",
        "奖罚款": "项目商务部 / 财务共享",
        "取费": "项目商务部 / 财务共享",
        "跑冒滴漏": "项目成本部 / 供应链管理",
    }
    return owner_map.get(category, "项目商务部")


def _remediation_actions(category: str) -> list[str]:
    action_map = {
        "工程量": ["复核竣工图、现场收方和结算清单", "按实际完成量办理核减或补证", "形成清单扣减记录"],
        "单价": ["核对合同单价和调价条款", "补充审批单、报价或市场依据", "确认未获批准价差核减金额"],
        "实际完成量": ["组织现场核验或资料复验", "补齐移交、验收、拆除等证据", "证据不足部分不予计入结算"],
        "变更签证": ["补齐授权审批链", "复核变更与原合同清单扣减关系", "未闭合前冻结对应结算支付"],
        "索赔": ["复核责任划分和工期影响", "补充费用测算与正式批复", "剔除未批准或重复索赔金额"],
        "奖罚款": ["勾稽奖罚台账、处罚单和结算扣减", "补扣已生效罚款", "未获批奖励不得计入结算"],
        "取费": ["复核合同费率、批准费率和计费基数", "核减超合同费率取费", "保留取费复算底稿"],
        "跑冒滴漏": ["勾稽材料、台班、日志和盘点记录", "追踪超耗超领责任", "核减无有效证据的成本"],
    }
    return action_map.get(category, ["补充审计证据", "明确责任部门和整改金额", "提交主审复核"])


def build_remediation_tasks(
    issues: list[dict[str, Any]],
    decisions: dict[str, Any] | None = None,
    updates: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Create close-out tasks for issues confirmed by the chief auditor."""
    decisions = decisions or {}
    updates = updates or {}
    tasks: list[dict[str, Any]] = []
    for issue in issues:
        decision = decisions.get(issue.get("id", ""), {})
        if decision.get("decision") != "confirm":
            continue
        update = updates.get(issue.get("id", ""), {})
        status = update.get("status") or "待整改"
        actions = _remediation_actions(issue.get("category", ""))
        task = {
            "task_id": f"rem-{issue.get('id', '')}",
            "issue_id": issue.get("id", ""),
            "title": issue.get("title", ""),
            "category": issue.get("category", ""),
            "risk": issue.get("risk", ""),
            "amount": issue.get("amount", 0),
            "owner": update.get("owner") or _remediation_owner(issue.get("category", "")),
            "status": status,
            "decision_time": decision.get("time", ""),
            "chief_note": decision.get("note", ""),
            "due_rule": "主审确认后 10 个工作日内完成整改反馈，20 个工作日内完成验证销号",
            "actions": actions,
            "required_evidence": [
                "整改说明或核减确认单",
                "责任部门复核底稿",
                "支撑资料页码、影像或系统截图",
                "主审验证记录",
            ],
            "latest_update": {
                "status": status,
                "note": update.get("note", ""),
                "time": update.get("time", ""),
                "operator": update.get("operator", ""),
            },
        }
        tasks.append(task)
    closed_statuses = {"已销号"}
    verified_statuses = {"整改验证", "已销号"}
    return {
        "workflow": ["主审确认", "下发整改", "责任部门反馈", "审计验证", "销号归档"],
        "tasks": tasks,
        "summary": {
            "total": len(tasks),
            "waiting": sum(1 for x in tasks if x["status"] == "待整改"),
            "in_progress": sum(1 for x in tasks if x["status"] == "整改中"),
            "verifying": sum(1 for x in tasks if x["status"] == "整改验证"),
            "closed": sum(1 for x in tasks if x["status"] in closed_statuses),
            "verified": sum(1 for x in tasks if x["status"] in verified_statuses),
            "amount": sum(x["amount"] for x in tasks),
            "open_amount": sum(x["amount"] for x in tasks if x["status"] not in closed_statuses),
        },
    }


def build_structured_model(
    data: dict[str, Any],
    comparisons: list[dict[str, Any]],
    issues: list[dict[str, Any]],
    cost_aggregation: dict[str, Any],
    fraud_assessment: dict[str, Any],
    agent_orchestration: dict[str, Any],
    audit_model_catalog: dict[str, Any],
    false_settlement_training: dict[str, Any],
    model_chain: dict[str, Any] | None = None,
    nine_grid: dict[str, Any] | None = None,
    cross_model_hints: dict[str, Any] | None = None,
) -> dict[str, Any]:
    project = data.get("project", {})
    documents = [
        {
            "id": doc.get("id", ""),
            "type": doc.get("type", ""),
            "name": doc.get("name", ""),
            "version": doc.get("version", ""),
            "location": doc.get("location", ""),
            "status": doc.get("status", ""),
            "fields": doc.get("fields", 0),
        }
        for doc in project.get("documents", [])
    ]
    items = []
    for row in comparisons:
        evidence = row.get("evidence", [])
        items.append(
            {
                "item_code": row.get("item_code", ""),
                "name": row.get("name", ""),
                "unit": row.get("unit", ""),
                "location": row.get("location", ""),
                "bid": row.get("bid", {}),
                "contract": row.get("contract", {}),
                "actual": row.get("actual", {}),
                "settlement": row.get("settlement", {}),
                "change": row.get("change", {}),
                "market_unit_price": row.get("market_unit_price"),
                "quantity_deviation_settlement_vs_actual": row.get("quantity_deviation_settlement_vs_actual"),
                "unit_price_deviation_settlement_vs_contract": row.get("unit_price_deviation_settlement_vs_contract"),
                "evidence_count": len(evidence),
                "evidence": evidence,
            }
        )
    lifecycle = deepcopy(data.get("lifecycle_findings", []))
    phase_models: list[dict[str, Any]] = []
    phase_index: dict[tuple[str, str], dict[str, Any]] = {}
    for item in lifecycle:
        key = (item.get("phase", ""), item.get("model_type", ""))
        slot = phase_index.get(key)
        if not slot:
            slot = {
                "phase": item.get("phase", ""),
                "model_type": item.get("model_type", ""),
                "domain": item.get("domain", ""),
                "count": 0,
                "amount": 0,
                "titles": [],
            }
            phase_index[key] = slot
            phase_models.append(slot)
        slot["count"] += 1
        slot["amount"] += _money(_number(item.get("amount")))
        slot["titles"].append(item.get("title", ""))
    return {
        "project": {
            "id": project.get("id", ""),
            "name": project.get("name", ""),
            "phase": project.get("phase", ""),
            "owner": project.get("owner", ""),
            "contractor": project.get("contractor", ""),
            "status": project.get("status", ""),
        },
        "documents": documents,
        "items": items,
        "cost_components": {
            "claims": deepcopy(project.get("cost_components", {}).get("claims", [])),
            "rewards_penalties": deepcopy(project.get("cost_components", {}).get("rewards_penalties", [])),
            "management_fee": deepcopy(project.get("cost_components", {}).get("management_fee", {})),
            "aggregation": deepcopy(cost_aggregation),
        },
        "lifecycle": lifecycle,
        "lifecycle_models": phase_models,
        "experts": deepcopy(data.get("experts", [])),
        "agents": deepcopy(agent_orchestration),
        "audit_model_catalog": deepcopy(audit_model_catalog),
        "false_settlement_training": deepcopy(false_settlement_training),
        "knowledge": deepcopy(data.get("knowledge", [])),
        "thresholds": deepcopy(data.get("meta", {}).get("thresholds", DEFAULT_THRESHOLDS)),
        "fraud_assessment": deepcopy(fraud_assessment),
        "model_chain": deepcopy(model_chain) if model_chain else None,
        "nine_grid": deepcopy(nine_grid) if nine_grid else None,
        "cross_model_hints": deepcopy(cross_model_hints) if cross_model_hints else None,
        "summary": {
            "document_count": len(documents),
            "item_count": len(items),
            "issue_count": len(issues),
            "historical_pattern_count": len(data.get("historical_patterns", [])),
            "historical_case_count": len(data.get("historical_cases", [])),
            "audit_model_count": audit_model_catalog.get("summary", {}).get("model_count", 0),
            "lifecycle_count": len(data.get("lifecycle_findings", [])),
        },
    }


def analyze(
    data: dict[str, Any],
    decisions: dict[str, Any] | None = None,
    remediation_updates: dict[str, Any] | None = None,
) -> dict[str, Any]:
    comparisons = compare_three_tables(data)
    issues = detect_issues(data, comparisons)
    issues.extend(detect_cost_component_issues(data))
    cost_aggregation = aggregate_cost_components(data)
    fraud_assessment = score_fraud_risk(data, issues)
    agent_orchestration = build_agent_orchestration(data, issues)
    core_results = run_core_models(data)
    model_chain = build_model_chain(core_results, issues)
    nine_grid = build_nine_grid(core_results, data)
    cross_model_hints = build_cross_model_hints(core_results, issues)
    summary = summarize(data, issues, comparisons, cost_aggregation)
    remediation_tasks = build_remediation_tasks(issues, decisions, remediation_updates)
    false_settlement_training = train_false_settlement_model(data, issues)
    audit_model_catalog = build_audit_model_catalog(data, issues, remediation_tasks)
    structured_model = build_structured_model(
        data,
        comparisons,
        issues,
        cost_aggregation,
        fraud_assessment,
        agent_orchestration,
        audit_model_catalog,
        false_settlement_training,
        model_chain=model_chain,
        nine_grid=nine_grid,
        cross_model_hints=cross_model_hints,
    )
    return {
        "meta": deepcopy(data.get("meta", {})),
        "organizations": deepcopy(data.get("organizations", [])),
        "project": deepcopy(data.get("project", {})),
        "comparisons": comparisons,
        "issues": issues,
        "summary": summary,
        "cost_aggregation": cost_aggregation,
        "fraud_assessment": fraud_assessment,
        "agent_orchestration": agent_orchestration,
        "audit_model_catalog": audit_model_catalog,
        "core_results": {mid: fs for mid, fs in core_results.items() if fs},
        "model_chain": model_chain,
        "nine_grid": nine_grid,
        "cross_model_hints": cross_model_hints,
        "false_settlement_training": false_settlement_training,
        "structured_model": structured_model,
        "experts": deepcopy(data.get("experts", [])),
        "expert_opinions": deepcopy(data.get("expert_opinions", {})),
        "lifecycle_findings": deepcopy(data.get("lifecycle_findings", [])),
        "knowledge": deepcopy(data.get("knowledge", [])),
        "closure": deepcopy(data.get("closure", [])),
        "remediation_tasks": remediation_tasks,
        "report": build_report(data, issues, summary, decisions, fraud_assessment, remediation_tasks),
    }


def answer_question(question: str, knowledge: list[dict[str, Any]], issue: dict[str, Any] | None = None) -> dict[str, Any]:
    query = (question or "").strip()
    if not query:
        return {"answer": "请输入制度、签证、结算、工程量或材料损耗问题。", "matches": []}
    ranked = []
    for entry in knowledge:
        score = sum(1 for word in entry.get("keywords", []) if word in query)
        if score:
            ranked.append((score, entry))
    ranked.sort(key=lambda pair: pair[0], reverse=True)
    matches = [entry for _, entry in ranked[:3]]
    if matches:
        primary = matches[0]
        answer = f"依据《{primary['title']}》{primary['basis']}：{primary['answer']}"
    else:
        answer = "当前知识库未命中明确条款。建议先核对合同专用条款、授权清单、变更审批链和实际完成量证据，并由主审人工确认。"
    if issue:
        answer += f" 当前关联疑点为“{issue.get('title', '')}”，涉及金额 {_money(_number(issue.get('amount'))):,} 元。"
    return {"answer": answer, "matches": matches}


def _tabular_rows_from_bytes(filename: str, payload: bytes) -> list[dict[str, Any]]:
    suffix = Path(filename).suffix.lower()
    if suffix == ".json":
        value = json.loads(payload.decode("utf-8-sig"))
        if isinstance(value, dict):
            return value.get("rows") or value.get("line_items") or [value]
        return value if isinstance(value, list) else []
    if suffix == ".csv":
        return list(csv.DictReader(io.StringIO(payload.decode("utf-8-sig"))))
    if suffix in {".xlsx", ".xls"}:
        try:
            import openpyxl
        except ImportError:
            return []
        workbook = openpyxl.load_workbook(io.BytesIO(payload), read_only=True, data_only=True)
        sheet = workbook.worksheets[0]
        values = list(sheet.values)
        if not values:
            return []
        headers = [str(x or "").strip() for x in values[0]]
        return [dict(zip(headers, row)) for row in values[1:] if any(x is not None for x in row)]
    return []


def _decode_text(payload: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    return payload.decode("utf-8", errors="ignore")


def structured_json_from_bytes(filename: str, payload: bytes) -> dict[str, Any] | None:
    if Path(filename).suffix.lower() != ".json":
        return None
    value = json.loads(payload.decode("utf-8-sig"))
    if isinstance(value, dict) and isinstance(value.get("line_items"), list):
        return value
    return None


def _preview_zip_member(name: str, payload: bytes) -> dict[str, Any]:
    preview = {
        "name": Path(name).name,
        "type": Path(name).suffix.lstrip("."),
        "document_type": classify_document_name(name),
        "fields": {},
        "rows": [],
        "warnings": [],
    }
    suffix = Path(name).suffix.lower()
    if suffix == ".json":
        structured = structured_json_from_bytes(name, payload)
        rows = _tabular_rows_from_bytes(name, payload)
        for row in rows:
            if isinstance(row, dict):
                row.setdefault("__document_name", Path(name).name)
                row.setdefault("__document_type", preview["document_type"])
        preview["rows"] = rows[:100]
        preview["fields"] = {"row_count": len(rows), "columns": list(rows[0].keys()) if rows else []}
        if structured:
            preview["structured_data"] = structured
            preview["fields"]["schema"] = "settlement_demo.line_items"
        return preview
    if suffix in {".csv", ".xlsx", ".xls"}:
        rows = _tabular_rows_from_bytes(name, payload)
        for row in rows:
            if isinstance(row, dict):
                row.setdefault("__document_name", Path(name).name)
                row.setdefault("__document_type", preview["document_type"])
        preview["rows"] = rows[:100]
        preview["fields"] = {"row_count": len(rows), "columns": list(rows[0].keys()) if rows else []}
        return preview
    if suffix in {".pdf", ".docx", ".txt"}:
        text, warnings, docx_rows = _extract_text_payload(name, payload)
        rows = docx_rows or _extract_text_rows(text, preview["document_type"], Path(name).name)
        for row in rows:
            row.setdefault("__document_name", Path(name).name)
            row.setdefault("__document_type", preview["document_type"])
        preview["rows"] = rows[:100]
        preview["fields"] = _extract_text_fields(text)
        preview["fields"].update({
            "text_length": len(text),
            "row_count": len(rows),
            "columns": list(rows[0].keys()) if rows else [],
        })
        preview["warnings"].extend(warnings)
        if not rows and suffix in {".pdf", ".docx"}:
            preview["warnings"].append("已抽取文本但未识别到标准清单行，可补充 OCR 或表头映射。")
        return preview
    preview["warnings"].append("Zip 内子文件类型暂未展开。")
    return preview


def classify_document_name(filename: str) -> str:
    name = Path(filename).stem
    categories = [
        ("合同", ("合同", "协议", "专用条款")),
        ("结算书", ("结算", "取费表")),
        ("变更签证", ("变更", "签证", "洽商")),
        ("竣工验收", ("竣工", "验收", "移交")),
        ("索赔资料", ("索赔", "工期索赔")),
        ("奖罚资料", ("奖罚", "奖励", "罚款", "处罚")),
        ("材料台账", ("材料", "收发存", "领用")),
        ("机械台班", ("机械", "台班", "施工日志")),
        ("招标文件", ("招标", "清单", "投标")),
    ]
    for label, words in categories:
        if any(word in name for word in words):
            return label
    return "未分类资料"


def _first(row: dict[str, Any], aliases: tuple[str, ...]) -> Any:
    lowered = {str(key).strip().lower(): value for key, value in row.items()}
    for alias in aliases:
        if alias in row:
            return row.get(alias)
        value = lowered.get(alias.lower())
        if value not in (None, ""):
            return value
    return None


def _infer_section(document_type: str, row: dict[str, Any]) -> str:
    explicit = str(_first(row, ("section", "source", "来源", "资料类型", "表类型", "阶段")) or "")
    text = f"{document_type} {explicit} {' '.join(str(key) for key in row.keys())}"
    if any(word in text for word in ("结算", "申报")):
        return "settlement"
    if any(word in text for word in ("合同", "协议")):
        return "contract"
    if any(word in text for word in ("实际", "完成", "竣工", "验收", "收方")):
        return "actual"
    if any(word in text for word in ("变更", "签证", "洽商")):
        return "change"
    if any(word in text for word in ("招标", "投标")):
        return "bid"
    return "settlement"


def normalize_uploaded_rows(preview: dict[str, Any]) -> list[dict[str, Any]]:
    normalized: list[dict[str, Any]] = []
    for index, row in enumerate(preview.get("rows", []), 1):
        if not isinstance(row, dict):
            continue
        document_type = str(row.get("__document_type") or preview.get("document_type", ""))
        document_name = str(row.get("__document_name") or preview.get("name", ""))
        item_code = str(_first(row, ("item_code", "项目编码", "清单编码", "项目编号", "编码")) or "").strip()
        name = str(_first(row, ("name", "项目名称", "清单子目", "子目名称", "工作内容")) or "").strip()
        if not item_code and not name:
            continue
        section = _infer_section(document_type, row)
        quantity = _number(_first(row, (
            "quantity",
            "工程量",
            "数量",
            "招标工程量",
            "合同工程量",
            "实际完成工程量",
            "结算工程量",
            "申报量",
        )))
        unit_price = _number(_first(row, (
            "unit_price",
            "单价",
            "综合单价",
            "招标单价",
            "合同单价",
            "结算单价",
            "市场参考价",
        )))
        amount = _number(_first(row, ("amount", "合价", "金额", "结算金额", "合同金额", "申报金额")))
        normalized.append({
            "item_code": item_code or f"UPLOAD-{index:03d}",
            "name": name or item_code,
            "unit": str(_first(row, ("unit", "单位", "计量单位")) or "").strip(),
            "location": str(_first(row, ("location", "部位", "工程部位", "施工部位")) or "").strip(),
            "section": section,
            "quantity": quantity,
            "unit_price": unit_price,
            "amount": amount or quantity * unit_price,
            "market_unit_price": _number(_first(row, ("market_unit_price", "市场参考价", "市场价"))),
            "document": document_name,
            "document_type": document_type,
            "source_row": index,
        })
    return normalized


def build_data_from_upload_previews(base_data: dict[str, Any], previews: list[dict[str, Any]]) -> tuple[dict[str, Any], bool]:
    """Build an active audit dataset from uploaded structured files or tables."""
    for preview in reversed(previews):
        structured = preview.get("structured_data")
        if isinstance(structured, dict) and isinstance(structured.get("line_items"), list):
            data = deepcopy(structured)
            data.setdefault("meta", {}).update({"data_source": "上传结构化资料包", "last_analysis": datetime_label()})
            return data, True

    records: list[dict[str, Any]] = []
    for preview in previews:
        records.extend(normalize_uploaded_rows(preview))
    if not records:
        return deepcopy(base_data), False

    data = deepcopy(base_data)
    data.setdefault("meta", {}).update({"data_source": "上传表格资料包", "last_analysis": datetime_label()})
    documents = []
    for index, preview in enumerate(previews, 1):
        documents.append({
            "id": f"upload-doc-{index:03d}",
            "type": preview.get("document_type", preview.get("type", "")),
            "name": preview.get("name", ""),
            "version": "上传版",
            "location": "上传资料包",
            "status": "已识别",
            "fields": preview.get("fields", {}).get("row_count", len(preview.get("rows", []))),
        })

    grouped: dict[str, dict[str, Any]] = {}
    for record in records:
        item = grouped.setdefault(record["item_code"], {
            "item_code": record["item_code"],
            "name": record["name"],
            "unit": record["unit"],
            "location": record["location"],
            "bid": {},
            "contract": {},
            "change": {"quantity": 0, "unit_price": 0, "amount": 0, "approved": True},
            "actual": {},
            "settlement": {},
            "evidence": [],
        })
        if record["name"]:
            item["name"] = record["name"]
        if record["unit"]:
            item["unit"] = record["unit"]
        if record["location"]:
            item["location"] = record["location"]
        section = record["section"]
        item[section] = {
            "quantity": record["quantity"],
            "unit_price": record["unit_price"],
            "amount": record["amount"],
        }
        if section == "change":
            item[section]["approved"] = True
        if record.get("market_unit_price"):
            item["market_unit_price"] = record["market_unit_price"]
        item["evidence"].append({
            "document": record["document"],
            "location": f"{record['document_type']} / 第 {record['source_row']} 行",
            "fact": f"{record['name']} {record['quantity']} {record['unit']}，单价 {record['unit_price']}，金额 {record['amount']}",
        })

    line_items = list(grouped.values())
    data["project"]["documents"] = documents
    data["line_items"] = line_items
    data["project"]["cost_components"] = {"claims": [], "rewards_penalties": [], "management_fee": {}}
    financials = data["project"].setdefault("financials", {})
    financials["contract_total"] = _money(sum(_amount(item.get("contract", {})) for item in line_items))
    financials["approved_change_total"] = _money(sum(_amount(item.get("change", {})) for item in line_items))
    financials["settlement_total"] = _money(sum(_amount(item.get("settlement", {})) for item in line_items))
    financials["actual_certified_total"] = _money(sum(_amount(item.get("actual", {})) for item in line_items))
    financials["paid_total"] = financials.get("paid_total", financials["actual_certified_total"])
    return data, True


def datetime_label() -> str:
    from datetime import datetime
    return datetime.now().strftime("%Y-%m-%d %H:%M")


def _extract_text_fields(text: str) -> dict[str, Any]:
    patterns = {
        "project_code": r"(?:合同编号|项目编号|项目编码)\s*[:：]?\s*([A-Za-z0-9-]+)",
        "quantity": r"(?:工程量|数量|申报量)\s*[:：]?\s*([\d,.]+)",
        "unit_price": r"(?:单价|含税单价|结算单价)\s*[:：]?\s*([\d,.]+)",
        "amount": r"(?:合价|金额|结算金额)\s*[:：]?\s*([\d,.]+)",
    }
    fields: dict[str, Any] = {}
    for name, pattern in patterns.items():
        match = re.search(pattern, text or "")
        if match:
            fields[name] = _number(match.group(1)) if name != "project_code" else match.group(1)
    return fields


def _extract_text_rows(text: str, document_type: str, document_name: str) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    line_pattern = re.compile(
        r"(?P<code>[A-Za-z0-9]{1,8}[-]?\d{2,}|[A-Za-z]{1,4}-\d{2,})"
        r"[\s,，|；;]+(?P<name>[\u4e00-\u9fa5A-Za-z0-9（）()·\- ]{2,40}?)"
        r"[\s,，|；;]+(?P<unit>m3|m³|㎡|m2|t|台班|项|米|吨|立方米|平方米)"
        r"[\s,，|；;]+(?P<quantity>-?\d+(?:,\d{3})*(?:\.\d+)?)"
        r"[\s,，|；;]+(?P<unit_price>-?\d+(?:,\d{3})*(?:\.\d+)?)"
        r"(?:[\s,，|；;]+(?P<amount>-?\d+(?:,\d{3})*(?:\.\d+)?))?"
    )
    labeled_pattern = re.compile(
        r"(?:项目编码|清单编码|编码)[:：]?\s*(?P<code>[A-Za-z0-9-]+).*?"
        r"(?:项目名称|清单子目|名称)[:：]?\s*(?P<name>[\u4e00-\u9fa5A-Za-z0-9（）()·\- ]+?)\s+"
        r"(?:单位|计量单位)[:：]?\s*(?P<unit>m3|m³|㎡|m2|t|台班|项|米|吨|立方米|平方米).*?"
        r"(?:工程量|数量|申报量)[:：]?\s*(?P<quantity>-?\d+(?:,\d{3})*(?:\.\d+)?).*?"
        r"(?:单价|综合单价|含税单价)[:：]?\s*(?P<unit_price>-?\d+(?:,\d{3})*(?:\.\d+)?)"
        r"(?:.*?(?:合价|金额|结算金额)[:：]?\s*(?P<amount>-?\d+(?:,\d{3})*(?:\.\d+)?))?"
    )
    for index, raw_line in enumerate((text or "").splitlines(), 1):
        line = re.sub(r"\s+", " ", raw_line.strip())
        if not line:
            continue
        match = labeled_pattern.search(line) or line_pattern.search(line)
        if not match:
            continue
        values = match.groupdict()
        quantity = _number(values.get("quantity"))
        unit_price = _number(values.get("unit_price"))
        amount = _number(values.get("amount"))
        rows.append({
            "项目编码": str(values.get("code") or "").strip(),
            "项目名称": str(values.get("name") or "").strip(" ：:，,;；"),
            "单位": str(values.get("unit") or "").strip(),
            "工程量": quantity,
            "单价": unit_price,
            "合价": amount or quantity * unit_price,
            "资料类型": document_type,
            "来源位置": f"{document_name} / 文本第 {index} 行",
            "__document_name": document_name,
            "__document_type": document_type,
        })
    return rows


def _extract_rows_from_docx_tables(document: Any, document_type: str, document_name: str) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables, 1):
        table_rows = table.rows
        if not table_rows:
            continue
        headers = [cell.text.strip() for cell in table_rows[0].cells]
        if not any(header for header in headers):
            continue
        for row_index, row in enumerate(table_rows[1:], 1):
            values = [cell.text.strip() for cell in row.cells]
            record = dict(zip(headers, values))
            if not (_first(record, ("项目编码", "清单编码", "编码", "item_code")) or _first(record, ("项目名称", "清单子目", "名称", "name"))):
                continue
            record.setdefault("资料类型", document_type)
            record.setdefault("__document_name", document_name)
            record.setdefault("__document_type", document_type)
            record.setdefault("来源位置", f"{document_name} / 表 {table_index} 第 {row_index} 行")
            rows.append(record)
    return rows


def _extract_text_payload(filename: str, payload: bytes) -> tuple[str, list[str], list[dict[str, Any]]]:
    suffix = Path(filename).suffix.lower()
    warnings: list[str] = []
    table_rows: list[dict[str, Any]] = []
    if suffix == ".txt":
        return _decode_text(payload), warnings, table_rows
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
            text = "\n".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(payload)).pages)
            return text, warnings, table_rows
        except ImportError:
            warnings.append("当前运行环境未安装 PDF 文本提取依赖，已接收文件但未抽取字段。")
        except Exception as exc:
            warnings.append(f"PDF 解析失败：{exc}")
        return "", warnings, table_rows
    if suffix == ".docx":
        try:
            from docx import Document
            document = Document(io.BytesIO(payload))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            table_rows = _extract_rows_from_docx_tables(document, classify_document_name(filename), Path(filename).name)
            return text, warnings, table_rows
        except ImportError:
            warnings.append("当前运行环境未安装 DOCX 文本提取依赖，已接收文件但未抽取字段。")
        except Exception as exc:
            warnings.append(f"DOCX 解析失败：{exc}")
    return "", warnings, table_rows


def parse_uploaded_file(path: str | Path) -> dict[str, Any]:
    """Extract a small, auditable preview from common settlement file types."""
    file_path = Path(path)
    payload = file_path.read_bytes()
    suffix = file_path.suffix.lower()
    result = {
        "name": file_path.name,
        "type": suffix.lstrip("."),
        "document_type": classify_document_name(file_path.name),
        "fields": {},
        "rows": [],
        "warnings": [],
    }
    if suffix in {".json", ".csv", ".xlsx", ".xls"}:
        structured = structured_json_from_bytes(file_path.name, payload) if suffix == ".json" else None
        rows = _tabular_rows_from_bytes(file_path.name, payload)
        result["rows"] = rows[:100]
        result["fields"] = {"row_count": len(rows), "columns": list(rows[0].keys()) if rows else []}
        if structured:
            result["structured_data"] = structured
            result["fields"]["schema"] = "settlement_demo.line_items"
        if not rows:
            result["warnings"].append("未识别到可结构化行，请检查表头或文件格式。")
        return result
    text = ""
    if suffix in {".pdf", ".docx", ".txt"}:
        text, warnings, table_rows = _extract_text_payload(file_path.name, payload)
        rows = table_rows or _extract_text_rows(text, result["document_type"], file_path.name)
        result["rows"] = rows[:100]
        result["fields"] = _extract_text_fields(text)
        result["fields"].update({
            "text_length": len(text),
            "row_count": len(rows),
            "columns": list(rows[0].keys()) if rows else [],
        })
        result["warnings"].extend(warnings)
        if not text:
            result["warnings"].append("未提取到文本字段，扫描件需要 OCR 适配后再入库。")
        elif not rows:
            result["warnings"].append("已抽取文本但未识别到标准清单行，可补充 OCR 或表头映射。")
        return result
    if suffix == ".zip":
        with zipfile.ZipFile(io.BytesIO(payload)) as archive:
            children = [name for name in archive.namelist() if not name.endswith("/")]
            child_previews = []
            merged_rows: list[dict[str, Any]] = []
            structured_data = None
            for child_name in children[:30]:
                try:
                    child_bytes = archive.read(child_name)
                except Exception:
                    continue
                child_preview = _preview_zip_member(child_name, child_bytes)
                child_previews.append(child_preview)
                merged_rows.extend(child_preview.get("rows", []))
                if structured_data is None and child_preview.get("structured_data"):
                    structured_data = child_preview["structured_data"]
        result["fields"] = {"file_count": len(children), "children": children[:30], "row_count": len(merged_rows)}
        result["rows"] = merged_rows[:100]
        result["child_previews"] = child_previews
        if structured_data:
            result["structured_data"] = structured_data
            result["fields"]["schema"] = "settlement_demo.line_items"
        result["warnings"].append("资料包已接收，子文件已按类型展开预览。")
        return result
    else:
        result["warnings"].append("文件已接收，当前 Demo 仅对 PDF、DOCX、Excel、CSV、JSON 和 ZIP 做字段预览。")
    return result
